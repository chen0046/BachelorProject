# import os
# import docx
# import PyPDF2

# def extract_text_from_docx(file_path):
#     doc = docx.Document(file_path)
#     text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
#     return text

# def extract_text_from_pdf(file_path):
#     with open(file_path, 'rb') as file:
#         pdf_reader = PyPDF2.PdfReader(file)
#         text = ''
#         for page_num in range(len(pdf_reader.pages)):
#             text += pdf_reader.pages[page_num].extract_text()
#     return text

# def find_keyword_in_documents_folder(folder_path, keyword):
#     # Initialize a variable to keep track of occurrences across all documents
#     total_keyword_count = 0

#     # Iterate through all files in the folder
#     for filename in os.listdir(folder_path):
#         if filename.endswith(('.docx', '.pdf', '.txt')):
#             # Construct the full path for each document
#             file_path = os.path.join(folder_path, filename)

#             # Extract text from the document
#             if filename.endswith('.docx'):
#                 document_text = extract_text_from_docx(file_path)
#             elif filename.endswith('.pdf'):
#                 document_text = extract_text_from_pdf(file_path)
#             elif filename.endswith('.txt'):
#                 with open(file_path, 'r', encoding='utf-8') as txt_file:
#                     document_text = txt_file.read()

#             # Initialize a variable to keep track of occurrences in the current document
#             keyword_count = document_text.lower().count(keyword.lower())

#             # Print the occurrences for the current document
#             print(f"Total occurrences of '{keyword}' in {filename}: {keyword_count}")

#             # Update the total count
#             total_keyword_count += keyword_count

#     # Print the total number of occurrences across all documents
#     print(f"Total occurrences of '{keyword}' in the entire folder: {total_keyword_count}")

# # Example usage with a folder path
# folder_path = '/Users/alexsolomon/Desktop/Bachelor/BachelorProject/Textract'
# keyword = 'medicine'
# find_keyword_in_documents_folder(folder_path, keyword)



####################################################################################
############### not using textract, -- finding files within folders ################
####################################################################################
######## pip install python-docx PyPDF2 

# import os
# import docx
# import PyPDF2

# def extract_text_from_docx(file_path):
#     doc = docx.Document(file_path)
#     text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
#     return text

# def extract_text_from_pdf(file_path):
#     with open(file_path, 'rb') as file:
#         pdf_reader = PyPDF2.PdfReader(file)
#         text = ''
#         for page_num in range(len(pdf_reader.pages)):
#             text += pdf_reader.pages[page_num].extract_text()
#     return text

# def find_keyword_in_documents_folder(folder_path, keyword):
#     # Initialize a variable to keep track of occurrences across all documents
#     total_keyword_count = 0

#     # Iterate through all files and subdirectories in the folder
#     for root, dirs, files in os.walk(folder_path):
#         for filename in files:
#             if filename.endswith(('.docx', '.pdf', '.txt')):
#                 # Construct the full path for each document
#                 file_path = os.path.join(root, filename)

#                 # Extract text from the document
#                 if filename.endswith('.docx'):
#                     document_text = extract_text_from_docx(file_path)
#                 elif filename.endswith('.pdf'):
#                     document_text = extract_text_from_pdf(file_path)
#                 elif filename.endswith('.txt'):
#                     with open(file_path, 'r', encoding='utf-8') as txt_file:
#                         document_text = txt_file.read()

#                 # Initialize a variable to keep track of occurrences in the current document
#                 keyword_count = document_text.lower().count(keyword.lower())

#                 # Print the occurrences for the current document
#                 print(f"Total occurrences of '{keyword}' in {file_path}: {keyword_count}")

#                 # Update the total count
#                 total_keyword_count += keyword_count

#     # Print the total number of occurrences across all documents
#     print(f"Total occurrences of '{keyword}' in the entire folder: {total_keyword_count}")

# # Example usage with a folder path
# folder_path = '/Users/alexsolomon/Desktop/Bachelor/BachelorProject/Textract'
# keyword = 'medicine'
# find_keyword_in_documents_folder(folder_path, keyword)



################## more keywords ##############################
# import os
# import docx
# import PyPDF2

# def extract_text_from_docx(file_path):
#     doc = docx.Document(file_path)
#     text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
#     return text

# def extract_text_from_pdf(file_path):
#     with open(file_path, 'rb') as file:
#         pdf_reader = PyPDF2.PdfReader(file)
#         text = ''
#         for page_num in range(len(pdf_reader.pages)):
#             text += pdf_reader.pages[page_num].extract_text()
#     return text

# def find_keyword_in_documents_folder(folder_path, keywords):
#     # Initialize a dictionary to keep track of occurrences for each keyword
#     total_keyword_counts = {keyword: 0 for keyword in keywords}

#     # Iterate through all files and subdirectories in the folder
#     for root, dirs, files in os.walk(folder_path):
#         for filename in files:
#             if filename.endswith(('.docx', '.pdf', '.txt')):
#                 # Construct the full path for each document
#                 file_path = os.path.join(root, filename)

#                 # Extract text from the document
#                 if filename.endswith('.docx'):
#                     document_text = extract_text_from_docx(file_path)
#                 elif filename.endswith('.pdf'):
#                     document_text = extract_text_from_pdf(file_path)
#                 elif filename.endswith('.txt'):
#                     with open(file_path, 'r', encoding='utf-8') as txt_file:
#                         document_text = txt_file.read()

#                 # Iterate over each keyword
#                 for keyword in keywords:
#                     # Initialize a variable to keep track of occurrences in the current document for the current keyword
#                     keyword_count = document_text.lower().count(keyword.lower())

#                     # Print the occurrences for the current document and keyword
#                     print(f"Total occurrences of '{keyword}' in {file_path}: {keyword_count}")

#                     # Update the total count for the current keyword
#                     total_keyword_counts[keyword] += keyword_count

#     # Print the total number of occurrences across all documents for each keyword
#     for keyword, count in total_keyword_counts.items():
#         print(f"Total occurrences of '{keyword}' in the entire folder: {count}")

# # Example usage with a folder path and multiple keywords
# folder_path = '/Users/alexsolomon/Desktop/Bachelor/BachelorProject/Textract'
# keywords = ['medicine', 'time', 'ball']
# find_keyword_in_documents_folder(folder_path, keywords)


##################################################################
################## works with excel ##############################
##################################################################



import os
import docx
import PyPDF2
import openpyxl

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return text

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ''
        for page_num in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_num].extract_text()
    return text

def extract_text_from_excel(file_path):
    wb = openpyxl.load_workbook(file_path)
    text = ''
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row in sheet.iter_rows(values_only=True):
            text += ' '.join(str(cell) for cell in row) + '\n'
    return text

def find_keyword_in_documents_folder(folder_path, keywords):
    # Initialize a dictionary to keep track of occurrences for each keyword
    total_keyword_counts = {keyword: 0 for keyword in keywords}

    # Iterate through all files and subdirectories in the folder
    for root, dirs, files in os.walk(folder_path):
        for filename in files:
            if filename.endswith(('.docx', '.pdf', '.txt', '.xlsx')):
                # Construct the full path for each document
                file_path = os.path.join(root, filename)

                # Extract text from the document
                if filename.endswith('.docx'):
                    document_text = extract_text_from_docx(file_path)
                elif filename.endswith('.pdf'):
                    document_text = extract_text_from_pdf(file_path)
                elif filename.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as txt_file:
                        document_text = txt_file.read()
                elif filename.endswith('.xlsx'):
                    document_text = extract_text_from_excel(file_path)

                # Iterate over each keyword
                for keyword in keywords:
                    # Initialize a variable to keep track of occurrences in the current document for the current keyword
                    keyword_count = document_text.lower().count(keyword.lower())

                    # Print the occurrences for the current document and keyword
                    print(f"Total occurrences of '{keyword}' in {file_path}: {keyword_count}")

                    # Update the total count for the current keyword
                    total_keyword_counts[keyword] += keyword_count

    # Print the total number of occurrences across all documents for each keyword
    for keyword, count in total_keyword_counts.items():
        print(f"Total occurrences of '{keyword}' in the entire folder: {count}")

# Example usage with a folder path and multiple keywords
folder_path = '/Users/alexsolomon/Desktop/Bachelor/BachelorProject/Textract'
keywords = ['medicine', 'time', 'ball']
find_keyword_in_documents_folder(folder_path, keywords)
